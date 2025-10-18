import os
from pathlib import Path
import json
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas

DATA_DIR = Path("datafolder")

def ensure_dir() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)

def make_txt():
    (DATA_DIR / "sample.txt").write_text(
        "This is a sample TXT file for ingestion. It has multiple lines.\n" * 5,
        encoding="utf-8",
    )

def make_csv():
    df = pd.DataFrame({"name": ["Alice", "Bob"], "score": [95, 88]})
    df.to_csv(DATA_DIR / "sample.csv", index=False)

def make_json():
    data = {"title": "Sample JSON", "items": [{"id": 1, "v": "a"}, {"id": 2, "v": "b"}]}
    (DATA_DIR / "sample.json").write_text(json.dumps(data, indent=2), encoding="utf-8")

def make_xlsx():
    df = pd.DataFrame({"col1": [1, 2, 3], "col2": ["a", "b", "c"]})
    df.to_excel(DATA_DIR / "sample.xlsx", index=False)

def make_docx():
    doc = Document()
    doc.add_heading("Sample DOCX", 0)
    for _ in range(3):
        doc.add_paragraph("This is a paragraph in a DOCX file used for ingestion testing.")
    doc.save(DATA_DIR / "sample.docx")

def make_pdf():
    c = canvas.Canvas(str(DATA_DIR / "sample.pdf"))
    c.setFont("Helvetica", 12)
    for i in range(1, 6):
        c.drawString(72, 800 - i * 20, f"This is line {i} in a generated PDF for ingestion testing.")
    c.save()

def main():
    ensure_dir()
    make_txt(); make_csv(); make_json(); make_xlsx(); make_docx(); make_pdf()
    print(f"Generated samples in {DATA_DIR.resolve()}")

if __name__ == "__main__":
    main()


